import pyodbc
import pandas as pd
from docx import Document
from docx.shared import Pt
import os
import datetime
from tqdm import tqdm

# Database connection details (Use environment variables for security)
server = os.getenv('DB_SERVER', 'your_server')
database = os.getenv('DB_NAME', 'your_database')
username = os.getenv('DB_USER', 'your_username')
password = os.getenv('DB_PASS', 'your_password')

# Establishing connection with error handling
conn_str = f"DRIVER={{ODBC Driver 18 for SQL Server}};SERVER={server};DATABASE={database};UID={username};PWD={password}"


try:
    with pyodbc.connect(conn_str) as conn:
        # Query to get schema details
        query = """
        SELECT TABLE_NAME, COLUMN_NAME, DATA_TYPE, CHARACTER_MAXIMUM_LENGTH, IS_NULLABLE, COLUMN_DEFAULT
        FROM INFORMATION_SCHEMA.COLUMNS
        WHERE TABLE_SCHEMA = 'dbo'
        ORDER BY TABLE_NAME, ORDINAL_POSITION
        """

        # Fetch data into Pandas DataFrame
        df = pd.read_sql(query, conn)

except pyodbc.Error as e:
    print(f"Database connection failed: {e}")
    exit()

# Create a new Word document
doc = Document()
doc.add_heading('Database Schema Documentation', level=1)

# Get unique table names
table_names = df["TABLE_NAME"].unique()

# Progress bar for table processing
for table_name in tqdm(table_names, desc="Processing Tables"):
    doc.add_heading(f'{table_name}', level=2)  # Table name as heading

    table_df = df[df["TABLE_NAME"] == table_name]
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    # Add header row
    hdr_cells = table.add_row().cells
    hdr_cells[0].text = "Column Name"
    hdr_cells[1].text = "Data Type"

    for cell in hdr_cells:
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(12)

    # Add column data with progress bar
    for _, row in tqdm(table_df.iterrows(), total=len(table_df), desc=f"Processing {table_name}", leave=False):
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["COLUMN_NAME"])
        row_cells[1].text = str(row["DATA_TYPE"])

        # Bold column names
        run = row_cells[0].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(11)

# Save the document with timestamp
timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
file_path = f"Database_Schema_{timestamp}.docx"
doc.save(file_path)

print(f"Document saved as {file_path}")